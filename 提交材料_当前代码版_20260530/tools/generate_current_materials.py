from __future__ import annotations

import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
MATERIAL_ROOT = ROOT / "提交材料_当前代码版_20260530"
UPDATE_DIR = MATERIAL_ROOT / "01_第二版_已更新"
SCHEMA_DIR = MATERIAL_ROOT / "02_当前运行数据库结构"
PPT_WORKSPACE = ROOT / "outputs" / "manual-current-materials" / "presentations" / "campus-market-current"
SLIDES_DIR = PPT_WORKSPACE / "slides"
PREVIEW_DIR = PPT_WORKSPACE / "preview"
LAYOUT_DIR = PPT_WORKSPACE / "layout"


FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
]


def font_path() -> str:
    for candidate in FONT_CANDIDATES:
        if candidate.exists():
            return str(candidate)
    return "arial.ttf"


FONT_PATH = font_path()


TABLES = [
    {
        "name": "users",
        "title": "用户 users",
        "fields": [
            ("PK", "id"),
            ("UQ", "username"),
            ("", "password_hash"),
            ("", "nickname, phone"),
            ("", "balance, created_at"),
        ],
    },
    {
        "name": "categories",
        "title": "分类 categories",
        "fields": [
            ("PK", "id"),
            ("UQ", "name"),
        ],
    },
    {
        "name": "products",
        "title": "商品 products",
        "fields": [
            ("PK", "id"),
            ("FK", "seller_id -> users.id"),
            ("FK", "category_id -> categories.id"),
            ("", "title, description, price"),
            ("", "condition_level"),
            ("", "image_url 兼容缓存"),
            ("", "status, view_count, created_at"),
        ],
    },
    {
        "name": "product_images",
        "title": "商品图片 product_images",
        "fields": [
            ("PK", "id"),
            ("FK", "product_id -> products.id"),
            ("", "image_url"),
            ("", "is_cover, sort_no"),
            ("", "created_at"),
        ],
    },
    {
        "name": "favorites",
        "title": "收藏 favorites",
        "fields": [
            ("PK", "id"),
            ("UQ", "user_id + product_id"),
            ("FK", "user_id -> users.id"),
            ("FK", "product_id -> products.id"),
            ("", "created_at"),
        ],
    },
    {
        "name": "orders",
        "title": "订单 orders",
        "fields": [
            ("PK", "id"),
            ("UQ", "order_no"),
            ("FK", "product_id -> products.id"),
            ("FK", "buyer_id -> users.id"),
            ("FK", "seller_id -> users.id"),
            ("", "amount, status, address"),
            ("", "created_at, paid_at, completed_at"),
        ],
    },
    {
        "name": "payments",
        "title": "支付 payments",
        "fields": [
            ("PK", "id"),
            ("FK", "order_id -> orders.id"),
            ("", "amount"),
            ("", "method: balance"),
            ("", "status, created_at"),
        ],
    },
]


def ensure_dirs() -> None:
    UPDATE_DIR.mkdir(parents=True, exist_ok=True)
    SCHEMA_DIR.mkdir(parents=True, exist_ok=True)
    SLIDES_DIR.mkdir(parents=True, exist_ok=True)
    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    LAYOUT_DIR.mkdir(parents=True, exist_ok=True)


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, fill="#375a64", width=5):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    import math

    angle = math.atan2(y2 - y1, x2 - x1)
    arrow_len = 18
    arrow_w = 10
    p1 = (
        x2 - arrow_len * math.cos(angle) + arrow_w * math.sin(angle),
        y2 - arrow_len * math.sin(angle) - arrow_w * math.cos(angle),
    )
    p2 = (
        x2 - arrow_len * math.cos(angle) - arrow_w * math.sin(angle),
        y2 - arrow_len * math.sin(angle) + arrow_w * math.cos(angle),
    )
    draw.polygon([end, p1, p2], fill=fill)


def draw_table_box(draw, x, y, w, table):
    header_h = 54
    row_h = 31
    h = header_h + row_h * len(table["fields"]) + 22
    draw.rounded_rectangle((x + 8, y + 9, x + w + 8, y + h + 9), radius=18, fill="#d7e4e0")
    draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill="#ffffff", outline="#17353d", width=4)
    draw.rounded_rectangle((x, y, x + w, y + header_h), radius=18, fill="#116b6f")
    draw.rectangle((x, y + header_h - 18, x + w, y + header_h), fill="#116b6f")

    title_font = ImageFont.truetype(FONT_PATH, 25)
    field_font = ImageFont.truetype(FONT_PATH, 19)
    tag_font = ImageFont.truetype(FONT_PATH, 16)
    draw.text((x + 20, y + 14), table["title"], font=title_font, fill="#ffffff")
    cy = y + header_h + 17
    for tag, field in table["fields"]:
        color = "#c73f53" if tag in {"PK", "UQ"} else "#12777b" if tag == "FK" else "#23343b"
        if tag:
            draw.text((x + 20, cy), tag, font=tag_font, fill=color)
            draw.text((x + 58, cy - 1), field, font=field_font, fill="#263940")
        else:
            draw.text((x + 20, cy - 1), field, font=field_font, fill="#263940")
        cy += row_h
    return (x, y, x + w, y + h)


def generate_er_png() -> Path:
    out = UPDATE_DIR / "校园二手交易平台ER模型图.png"
    img = Image.new("RGB", (2400, 1550), "#f4f7f2")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(FONT_PATH, 58)
    sub_font = ImageFont.truetype(FONT_PATH, 29)
    label_font = ImageFont.truetype(FONT_PATH, 20)

    draw.rectangle((0, 0, 2400, 165), fill="#e8f1ed")
    draw.text((690, 32), "校园二手交易平台数据库 ER 模型图", font=title_font, fill="#0c2f4a")
    draw.text((700, 103), "当前代码版：12 张运行表 · 多图/收藏/消息/提醒/管理员", font=sub_font, fill="#546774")

    boxes = {
        "users": draw_table_box(draw, 95, 230, 420, TABLES[0]),
        "categories": draw_table_box(draw, 985, 225, 420, TABLES[1]),
        "products": draw_table_box(draw, 720, 560, 610, TABLES[2]),
        "product_images": draw_table_box(draw, 95, 900, 520, TABLES[3]),
        "favorites": draw_table_box(draw, 720, 1035, 520, TABLES[4]),
        "orders": draw_table_box(draw, 1570, 530, 610, TABLES[5]),
        "payments": draw_table_box(draw, 1610, 1045, 560, TABLES[6]),
    }

    draw_arrow(draw, (515, 390), (720, 650))
    draw.text((540, 505), "1:N 发布商品", font=label_font, fill="#0c4252")
    draw_arrow(draw, (1195, 355), (1085, 560))
    draw.text((1210, 445), "1:N 分类", font=label_font, fill="#0c4252")
    draw_arrow(draw, (720, 835), (585, 900))
    draw.text((410, 835), "1:N 商品图片", font=label_font, fill="#0c4252")
    draw_arrow(draw, (955, 985), (955, 1035))
    draw.text((970, 997), "1:N 收藏记录", font=label_font, fill="#0c4252")
    draw_arrow(draw, (1330, 720), (1570, 700))
    draw.text((1380, 668), "1:N 形成订单", font=label_font, fill="#0c4252")
    draw.text((1590, 485), "orders.buyer_id / seller_id 均关联 users.id", font=label_font, fill="#0c4252")
    draw_arrow(draw, (1900, 940), (1900, 1045))
    draw.text((1920, 980), "1:N 支付流水", font=label_font, fill="#0c4252")
    draw_arrow(draw, (515, 465), (720, 1095))
    draw.text((548, 920), "1:N 用户收藏", font=label_font, fill="#0c4252")

    note = "设计说明：products.image_url 仍保留为兼容缓存；正式多图关系由 product_images 承载，favorites 支撑我的收藏页面。"
    draw.rounded_rectangle((305, 1455, 2095, 1518), radius=18, fill="#fff7df", outline="#d9a12f", width=2)
    draw.text((340, 1472), note, font=label_font, fill="#684c12")
    img.save(out, "PNG")
    return out


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "222222") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True, "0B2545")
        set_cell_shading(table.rows[0].cells[i], "E8F1ED")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    document.add_paragraph()
    return table


def apply_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(11, 37, 69)
        style.font.bold = True


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(88, 100, 112)
    doc.add_paragraph()


def generate_report(er_path: Path) -> Path:
    out = UPDATE_DIR / "校园二手交易平台数据库课程设计报告.docx"
    doc = Document()
    apply_doc_styles(doc)
    add_title(doc, "校园二手交易平台数据库课程设计报告", "当前代码版 · Flask + MySQL 5.6 · 12 张运行表")

    add_table(
        doc,
        ["项目", "内容"],
        [
            ["课程名称", "数据库原理与应用"],
            ["项目主题", "校园二手交易平台"],
            ["数据库名", "secondhand"],
            ["运行环境", "Python Flask、MySQL 5.6、PyMySQL、Jinja2、原生 CSS"],
            ["核心规模", "12 张运行表、7 个基础分类、3 个演示账号、15 件真实感演示商品"],
            ["当前重点功能", "商品多图表、我的收藏、本地多图上传、商品编辑/删除/重新上架、余额支付、个人资料与密码修改"],
        ],
        [3.2, 12.4],
    )

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(
        "本课程设计实现了一个面向校园闲置物品流转的二手交易平台。当前代码版以 Flask Web 应用为核心，"
        "数据库使用 MySQL 5.6 和 InnoDB 引擎，覆盖用户注册登录、商品发布与浏览、本地多图上传、收藏商品、"
        "订单创建、余额支付、确认收货和卖家商品管理等流程。"
    )
    doc.add_paragraph(
        "当前可运行网站采用 12 张运行表：users、admins、categories、products、product_images、"
        "favorites、orders、payments。product_images 保存一个商品的多张图片并支持封面排序；favorites 保存用户收藏关系。"
        "products.image_url 仍保留为兼容缓存，便于旧数据和旧模板平滑过渡。"
    )

    doc.add_heading("1 项目背景与功能范围", level=1)
    add_table(
        doc,
        ["模块", "当前代码版功能"],
        [
            ["用户管理", "注册、登录、退出，个人中心可修改昵称、手机号和登录密码。"],
            ["商品浏览", "首页展示商品、支持关键词和分类筛选，分类包含“其他”。"],
            ["商品发布", "支持本地上传最多 4 张图片，也保留图片链接兜底；上传文件保存在 static/uploads。"],
            ["商品图片", "product_images 独立保存图片路径、封面标记和排序，详情页可切换多张图片。"],
            ["我的收藏", "买家可收藏/取消收藏商品，个人中心可进入我的收藏列表。"],
            ["卖家管理", "卖家可编辑商品信息和图片、删除未成交商品、下架后重新上架。"],
            ["交易流程", "买家下单后商品锁定，余额支付后进入待收货，确认收货后卖家收到款项。"],
        ],
        [3.0, 12.6],
    )

    doc.add_heading("2 概念结构设计与 ER 模型", level=1)
    doc.add_paragraph(
        "当前系统抽象为用户、分类、商品、商品图片、收藏、订单和支付七类实体。用户可以发布商品，也可以收藏或购买其他用户的商品；"
        "商品属于一个分类并拥有多张图片；订单关联商品、买家和卖家；支付流水关联订单。"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(er_path), width=Inches(6.8))
    caption = doc.add_paragraph("图 1 当前代码版 ER 模型图")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3 逻辑结构设计", level=1)
    add_table(
        doc,
        ["表名", "中文含义", "作用"],
        [
            ["users", "用户表", "保存登录账号、密码摘要、昵称、手机号和钱包余额。"],
            ["categories", "分类表", "维护商品分类，包括数码电子、图书教材、服饰鞋包、运动户外、生活家居、美妆个护、其他。"],
            ["products", "商品表", "保存卖家、分类、标题、描述、价格、成色、兼容图片字段、状态和浏览量。"],
            ["product_images", "商品图片表", "一个商品对应多张图片，记录图片路径、封面标记和排序。"],
            ["favorites", "收藏表", "一个用户可收藏多个商品，同一用户对同一商品只能收藏一次。"],
            ["orders", "订单表", "保存订单号、商品、买家、卖家、成交金额、地址和订单状态。"],
            ["payments", "支付流水表", "保存订单支付金额、支付方式和支付结果。"],
        ],
        [2.8, 3.2, 9.4],
    )

    doc.add_heading("3.1 关键字段说明", level=2)
    add_table(
        doc,
        ["对象", "关键字段", "设计说明"],
        [
            ["users", "username, password_hash, nickname, phone, balance", "username 唯一；密码只存摘要；balance 用于余额支付演示。"],
            ["products", "seller_id, category_id, image_url, status", "seller_id 和 category_id 为外键；image_url 为兼容缓存；status 控制商品生命周期。"],
            ["product_images", "product_id, image_url, is_cover, sort_no", "支持一个商品多张图，删除商品时图片记录级联删除。"],
            ["favorites", "user_id, product_id", "设置 user_id + product_id 唯一键，避免重复收藏。"],
            ["orders", "order_no, product_id, buyer_id, seller_id, amount, status", "订单同时记录买卖双方，方便买家订单和卖家订单页面查询。"],
            ["payments", "order_id, amount, method, status", "当前支付方式为 balance，不保存银行卡等敏感信息。"],
        ],
        [2.6, 5.1, 7.7],
    )

    doc.add_heading("4 图片上传与文件路径设计", level=1)
    add_table(
        doc,
        ["项目", "当前设计"],
        [
            ["上传数量", "每个商品最多保留 4 张图片。"],
            ["上传大小", "应用级总上传限制为 20MB。"],
            ["用户上传路径", "代码/static/uploads/，该目录属于运行产生的数据，不随代码备份。"],
            ["演示图片路径", "代码/static/product_images/，用于初始化演示商品。"],
            ["数据库保存方式", "product_images.image_url 保存每张 /static/... 路径；products.image_url 同步保存兼容缓存。"],
        ],
        [3.2, 12.4],
    )

    doc.add_heading("5 SQL 实施与初始化数据", level=1)
    doc.add_paragraph(
        "当前网站实际使用的建表文件为代码/database/schema.sql，初始化入口为代码/init_db.py。执行 python init_db.py 会重建 secondhand 数据库，"
        "创建七张表、写入基础分类、演示账号、商品数据和对应商品图片记录。"
    )
    add_table(
        doc,
        ["账号", "密码", "用途"],
        [
            ["alice", "123456", "演示卖家/买家，初始余额 5000 元。"],
            ["bob", "123456", "演示卖家/买家，初始余额 5000 元。"],
            ["carol", "123456", "演示卖家/买家，初始余额 5000 元。"],
        ],
        [3.0, 3.0, 9.6],
    )

    doc.add_heading("6 完整性与安全设计", level=1)
    add_table(
        doc,
        ["风险", "当前处理方式"],
        [
            ["重复账号", "users.username 设置唯一键。"],
            ["重复收藏", "favorites.user_id + product_id 设置唯一键。"],
            ["孤立商品图片或收藏", "product_images.product_id、favorites.product_id 使用级联外键。"],
            ["订单引用无效商品或用户", "orders.product_id、buyer_id、seller_id 使用外键。"],
            ["支付流水脱离订单", "payments.order_id 使用外键。"],
            ["明文密码风险", "password_hash 保存密码摘要，不保存明文密码。"],
        ],
        [4.0, 11.6],
    )

    doc.add_heading("7 总结", level=1)
    doc.add_paragraph(
        "当前代码版从基础交易流程扩展为 12 表运行结构：商品多图、收藏、订单、支付、评价、消息、管理员权限和事件提醒都已落表。"
        "网站仍保留轻量、可运行、易演示的特点，同时更贴近真实交易平台。"
    )

    doc.save(out)
    return out


def generate_instruction() -> Path:
    out = UPDATE_DIR / "校园二手交易平台作业说明书.docx"
    doc = Document()
    apply_doc_styles(doc)
    add_title(doc, "作业提交与答辩说明书", "当前代码版材料说明")

    doc.add_heading("1 文件口径", level=1)
    doc.add_paragraph(
        "本文件夹中的当前代码版材料以实际可运行网站为准。当前运行库为 12 张表，"
        "旧版 10 表材料仍保存在 00_旧版重要参考_勿改 中，只作为扩展设计参考。"
    )
    add_table(
        doc,
        ["文件", "用途"],
        [
            ["校园二手交易平台数据库课程设计报告.docx", "当前代码版正式报告，说明 12 张运行表和网站功能。"],
            ["校园二手交易平台数据库课程设计答辩PPT.pptx", "8 页当前功能演示版 PPT。"],
            ["校园二手交易平台数据库课程设计答辩PPT_详细版.pptx", "12 页详细答辩版 PPT。"],
            ["校园二手交易平台ER模型图.png", "当前 12 表 ER 图或当前运行 schema，可插入报告或 PPT。"],
            ["../02_当前运行数据库结构/当前网站实际使用_schema.sql", "当前 Flask 网站实际执行的数据库结构文件。"],
        ],
        [7.0, 8.4],
    )

    doc.add_heading("2 答辩时怎么讲", level=1)
    for text in [
        "先讲现实背景：校园闲置物品流转频繁，微信群交易图片和状态管理不方便。",
        "再讲当前系统：注册登录、浏览筛选、发布多图商品、收藏商品、下单支付、确认收货、卖家管理和个人资料修改。",
        "展示 ER 图：强调当前实际运行 12 张表，其中 product_images 负责多图，notifications 负责事件提醒。",
        "讲交易状态：商品从 on_sale 到 locked，再到 sold；下架商品为 removed，可重新上架。",
        "讲图片路径：演示图在 static/product_images，用户上传图在 static/uploads，数据库保存 /static/... 访问路径。",
        "最后讲可维护性：评价、站内消息、管理员权限、事件提醒已经拆表，后续可继续扩展。",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("3 当前数据库结构", level=1)
    add_table(
        doc,
        ["表", "一句话说明"],
        [
            ["users", "账号、密码摘要、昵称、手机号、余额。"],
            ["categories", "商品分类，包含“其他”。"],
            ["products", "商品主体，含兼容图片字段、价格、状态。"],
            ["product_images", "商品多图路径、封面标记和排序。"],
            ["favorites", "用户收藏商品关系。"],
            ["orders", "买卖双方、商品、金额、地址、订单状态。"],
            ["payments", "余额支付流水。"],
        ],
        [4.0, 11.6],
    )

    doc.add_heading("4 提交前检查", level=1)
    for text in [
        "确认报告、PPT、ER 图都写的是 secondhand 数据库和 12 张运行表。",
        "不要把旧版 10 张业务表材料误当成当前代码版提交。",
        "如果老师要求演示网站，则以代码/database/schema.sql 和 init_db.py 为准。",
        "用户上传图片在 static/uploads，属于运行数据，不随 GitHub 代码备份。",
    ]:
        doc.add_paragraph(text)

    doc.save(out)
    return out


COMMON_JS = """
const C = {
  ink: "#0B2545", teal: "#116B6F", mint: "#E8F1ED", paper: "#F7F5EF",
  gold: "#C7912D", gray: "#5D6B78", line: "#D4DDD8", white: "#FFFFFF"
};
function bg(slide, ctx, title, kicker = "") {
  ctx.addShape(slide, { x: 0, y: 0, w: ctx.W, h: ctx.H, fill: C.paper });
  ctx.addShape(slide, { x: 0, y: 0, w: ctx.W, h: 86, fill: C.mint });
  ctx.addText(slide, { text: kicker, x: 58, y: 24, w: 220, h: 24, fontSize: 13, color: C.teal, bold: true });
  ctx.addText(slide, { text: title, x: 58, y: 48, w: 930, h: 38, fontSize: 26, color: C.ink, bold: true });
  ctx.addText(slide, { text: String(ctx.slideNumber).padStart(2, "0"), x: 1160, y: 650, w: 50, h: 24, fontSize: 12, color: C.gray, align: "right" });
}
function card(slide, ctx, x, y, w, h, title, text, accent = C.teal) {
  ctx.addShape(slide, { x, y, w, h, fill: C.white, line: ctx.line(C.line, 1) });
  ctx.addShape(slide, { x, y, w, h: 7, fill: accent });
  ctx.addText(slide, { text: title, x: x + 18, y: y + 18, w: w - 36, h: 26, fontSize: 18, color: C.ink, bold: true });
  ctx.addText(slide, { text, x: x + 18, y: y + 56, w: w - 36, h: h - 68, fontSize: 14, color: C.gray });
}
function bulletList(slide, ctx, items, x, y, w, size = 17) {
  items.forEach((item, i) => {
    const yy = y + i * 48;
    ctx.addShape(slide, { x, y: yy + 5, w: 12, h: 12, fill: i % 2 ? C.gold : C.teal });
    ctx.addText(slide, { text: item, x: x + 28, y: yy, w, h: 38, fontSize: size, color: C.ink });
  });
}
""".strip()


def js(value) -> str:
    return json.dumps(value, ensure_ascii=False)


def slide_module(slide_no: int, body: str) -> None:
    (SLIDES_DIR / f"slide-{slide_no:02d}.mjs").write_text(body, encoding="utf-8")


def generate_card_slide(no: int, title: str, kicker: str, cards, bullets=None) -> None:
    slide_module(no, f"""
{COMMON_JS}
export async function slide{no:02d}(presentation, ctx) {{
  const slide = presentation.slides.add();
  bg(slide, ctx, {js(title)}, {js(kicker)});
  const cards = {js(cards)};
  const positions = [[70, 135], [350, 135], [630, 135], [910, 135], [210, 340], [630, 340]];
  cards.forEach((c, i) => card(slide, ctx, positions[i][0], positions[i][1], i < 4 ? 250 : 360, i < 4 ? 160 : 145, c[0], c[1], i % 2 ? C.gold : C.teal));
  bulletList(slide, ctx, {js(bullets or [])}, 150, 545, 920, 16);
  return slide;
}}
""")


def generate_slide_modules(er_path: Path, detailed: bool = False) -> None:
    for file in SLIDES_DIR.glob("slide-*.mjs"):
        file.unlink()

    er_js = er_path.as_posix()
    slide_module(1, f"""
{COMMON_JS}
export async function slide01(presentation, ctx) {{
  const slide = presentation.slides.add();
  ctx.addShape(slide, {{ x: 0, y: 0, w: ctx.W, h: ctx.H, fill: "#0B2545" }});
  ctx.addShape(slide, {{ x: 0, y: 500, w: ctx.W, h: 220, fill: "#E8F1ED" }});
  ctx.addText(slide, {{ text: "校园二手交易平台", x: 72, y: 105, w: 760, h: 74, fontSize: 54, color: C.white, bold: true }});
  ctx.addText(slide, {{ text: "数据库设计与实现 · 第一轮扩展版", x: 76, y: 190, w: 720, h: 36, fontSize: 24, color: "#D9E9E5" }});
  card(slide, ctx, 76, 535, 260, 92, "12 张运行表", "多图/收藏/提醒/管理员", C.gold);
  card(slide, ctx, 370, 535, 280, 92, "独立多图表", "商品图片可封面排序，旧字段作缓存", C.teal);
  card(slide, ctx, 684, 535, 300, 92, "我的收藏", "详情页收藏，个人中心查看收藏列表", C.gold);
  await ctx.addImage(slide, {{ path: {js(er_js)}, x: 815, y: 82, w: 360, h: 260, fit: "cover", alt: "当前 ER 图缩略图" }});
  ctx.addText(slide, {{ text: "数据库名：secondhand · Flask + MySQL 5.6", x: 76, y: 650, w: 760, h: 24, fontSize: 14, color: C.ink }});
  return slide;
}}
""")

    generate_card_slide(2, "当前版本以可运行网站为准", "01 / 功能范围", [
        ["账号与个人中心", "注册登录、退出、修改昵称/手机号、修改密码。"],
        ["商品发布", "本地上传多张图片，编辑时可追加或删除旧图。"],
        ["商品管理", "编辑、删除、下架后重新上架。"],
        ["订单支付", "余额支付、平台托管、确认收货后打款。"],
        ["商品多图", "product_images 保存每张图片路径、封面标记和排序。"],
        ["我的收藏", "favorites 记录用户收藏关系，个人中心可查看。"],
    ], ["分类包含“其他”，演示商品更贴近校园场景", "上传目录和数据库路径口径已经在说明书中同步"])

    slide_module(3, f"""
{COMMON_JS}
export async function slide03(presentation, ctx) {{
  const slide = presentation.slides.add();
  bg(slide, ctx, "数据库结构从基础表扩展为当前 12 表", "02 / 数据库口径");
  const data = [
    ["users", "用户、密码摘要、昵称、余额"],
    ["categories", "商品分类，包含其他"],
    ["products", "商品主体、状态、兼容图片缓存"],
    ["product_images", "商品多图路径、封面与排序"],
    ["favorites", "用户收藏商品关系"],
    ["orders", "订单、买卖双方、金额、地址"],
    ["payments", "余额支付流水"]
  ];
  data.forEach((row, i) => {{
    const y = 126 + i * 68;
    ctx.addShape(slide, {{ x: 110, y, w: 225, h: 44, fill: i % 2 ? C.gold : C.teal }});
    ctx.addText(slide, {{ text: row[0], x: 130, y: y + 11, w: 185, h: 22, fontSize: 18, color: C.white, bold: true }});
    ctx.addText(slide, {{ text: row[1], x: 370, y: y + 9, w: 690, h: 28, fontSize: 21, color: C.ink }});
    ctx.addShape(slide, {{ x: 348, y: y + 22, w: 650, h: 2, fill: "#D8E0DC" }});
  }});
  ctx.addText(slide, {{ text: "当前报告、PPT、ER 图和 schema.sql 都按 7 表口径更新；旧版 10 表只作参考。", x: 120, y: 620, w: 930, h: 32, fontSize: 18, color: C.gray }});
  return slide;
}}
""")

    slide_module(4, f"""
{COMMON_JS}
export async function slide04(presentation, ctx) {{
  const slide = presentation.slides.add();
  bg(slide, ctx, "ER 图展示当前实际运行的表和外键", "03 / ER 模型");
  await ctx.addImage(slide, {{ path: {js(er_js)}, x: 54, y: 112, w: 1110, h: 540, fit: "contain", alt: "当前代码版 ER 图" }});
  ctx.addText(slide, {{ text: "重点：product_images 让多图关系规范化，favorites 支撑我的收藏；products.image_url 只保留兼容缓存。", x: 110, y: 650, w: 1000, h: 24, fontSize: 16, color: C.gray, align: "center" }});
  return slide;
}}
""")

    generate_card_slide(5, "图片上传和收藏是第一轮扩展重点", "04 / 功能扩展", [
        ["上传体验", "可一次选择多张，也可编辑时追加新图。"],
        ["存储路径", "用户上传：代码/static/uploads；演示图：static/product_images。"],
        ["图片表", "product_images 保存 /static/... 路径，支持封面和排序。"],
        ["收藏表", "favorites 用唯一键避免重复收藏。"],
    ], ["每个商品最多保留 4 张图片", "详情页可切换多图并显示收藏人数", "个人中心新增“我的收藏”入口"])

    generate_card_slide(6, "商品和订单状态让交易流程可控", "05 / 交易状态", [
        ["on_sale", "在售，可被购买和收藏。"],
        ["locked", "已下单锁定，避免重复购买。"],
        ["paid", "余额已扣，平台托管中。"],
        ["completed / sold", "确认收货后订单完成、商品售出。"],
    ], ["待支付订单取消后商品回到 on_sale", "卖家下架为 removed，后续可以重新上架"])

    generate_card_slide(7, "初始化数据服务演示，不再像测试占位", "06 / 演示数据", [
        ["演示账号", "alice / bob / carol，密码均为 123456。"],
        ["商品类型", "手机、教材、鼠标、自行车、鞋、台灯、宿舍折叠桌板。"],
        ["分类", "数码电子、图书教材、服饰鞋包、运动户外、生活家居、美妆个护、其他。"],
        ["验证脚本", "_verify.py 会检查 product_images 和 favorites。"],
    ], ["schema.sql 是当前网站实际使用的建表文件", "用户上传图片不随代码备份，避免把运行数据混进 GitHub"])

    generate_card_slide(8, "答辩收束：当前版本可运行、可演示、可继续扩展", "07 / 总结", [
        ["已经完成", "7 表数据库、独立商品图片表、收藏表、多图上传、余额交易。"],
        ["设计取舍", "保留 image_url 兼容缓存，避免旧数据和模板一次性断裂。"],
        ["后续扩展", "评价、消息、管理员权限和事件提醒已经落表。"],
        ["一句话", "以 secondhand 数据库和当前 schema.sql 为准。"],
    ])

    if detailed:
        generate_card_slide(9, "字段设计围绕页面查询和交易动作展开", "08 / 字段细节", [
            ["product_images", "product_id / image_url / is_cover / sort_no。"],
            ["favorites", "user_id + product_id 唯一，防止重复收藏。"],
            ["products.status", "on_sale / locked / sold / removed。"],
            ["orders.status", "created / paid / completed / cancelled。"],
        ], ["删除商品时 product_images 和 favorites 级联清理", "订单与支付保留独立流水，便于演示资金变化"])
        generate_card_slide(10, "页面功能和数据库字段一一对应", "09 / 页面到数据", [
            ["发布页", "写入 products，并同步 product_images。"],
            ["详情页", "读取 product_images 和 favorites，展示多图和收藏状态。"],
            ["个人中心", "展示发布、买到、卖出、收藏数量。"],
            ["订单页", "读取 orders 并关联 products、买家和卖家。"],
        ])
        generate_card_slide(11, "提交材料分层保留，避免旧版和当前版混用", "10 / 文件管理", [
            ["00_旧版重要参考_勿改", "保留旧版 10 表报告、PPT、ER 图和 SQL。"],
            ["01_第二版_已更新", "本次更新后的报告、PPT、ER 图，和当前代码一致。"],
            ["02_当前运行数据库结构", "存放当前网站真正使用的 schema.sql。"],
            ["GitHub 备份", "只备份代码，不把 static/uploads 运行图片放进去。"],
        ])
        generate_card_slide(12, "老师追问时的回答口径", "11 / 答辩问答", [
            ["为什么 12 张表？", "交易闭环加多图、收藏、消息、评价、管理员和事件提醒，功能与 ER 更完整。"],
            ["多图为什么还保留 image_url？", "它是兼容缓存，正式多图以 product_images 为准。"],
            ["旧版 10 表还能用吗？", "可以作为扩展设计参考，演示以当前 schema.sql 为准。"],
            ["上传图片在哪里？", "用户上传图在 代码/static/uploads，演示图在 static/product_images。"],
        ])


def generate_all(detailed: bool = True) -> dict:
    ensure_dirs()
    er = generate_er_png()
    report = generate_report(er)
    instruction = generate_instruction()
    generate_slide_modules(er, detailed=detailed)
    schema_src = ROOT / "代码" / "database" / "schema.sql"
    if schema_src.exists():
        (SCHEMA_DIR / "当前网站实际使用_schema.sql").write_bytes(schema_src.read_bytes())
    return {
        "er": str(er),
        "report": str(report),
        "instruction": str(instruction),
        "ppt_workspace": str(PPT_WORKSPACE),
        "slides_dir": str(SLIDES_DIR),
    }


if __name__ == "__main__":
    result = generate_all()
    print(json.dumps(result, ensure_ascii=False, indent=2))


